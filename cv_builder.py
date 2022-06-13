from docx import Document
from docx.shared import Inches
import pyttsx3

# pyttsx3.speak('Hello')

def speak(text):
    pyttsx3.speak(text)

# create and save document
document = Document()
# profile picture
document.add_picture('profile-pic.png', width=Inches(2.0))

# Inputting name ,phone number and email addredd
name = input('What is your name ?')
speak('Hello ' +name + "How are you today")

speak('what is your phone number ? ')
phone_number = input('what is your phone number ? ')
email = input('what is your email ? ')

# adding info 
document.add_paragraph(name + ' | ' + phone_number + ' | '+email)

#about me
document.add_heading('About me')
about_me = input("Tell about yourself ")
document.add_paragraph(about_me)

#work experience
document.add_heading("Work experience")
p = document.add_paragraph()

company = input('Enter company ')
from_date= input('From date ')
to_date = input('To Date ')

p.add_run(company + " " ).bold = True
p.add_run(from_date+ "-" + to_date +'\n').italic = True

experience_details = input('Describe your experience at ' + company + " ")

p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input('Do you have more experiences ? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date= input('From date ')
        to_date = input('To Date ')

        p.add_run(company + " " ).bold = True
        p.add_run(from_date+ "-" + to_date +'\n').italic = True

        experience_details = input('Describe your experience at ' + company)

        p.add_run(experience_details)
    else:
        break    

#Skills
document.add_heading('Skills')
skill = input('Enter skill ')
p =document.add_paragraph(skill) 
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills ? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill ')
        p =document.add_paragraph(skill) 
        p.style = 'List Bullet'
    else:
        break   


# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using ArBa Development Studios"





document.save('cv.docx')


# Download packages from requirements.txt
# pip3 install -r requirements.txt
